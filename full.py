
import google.generativeai as genai
import gradio as gr
import speech_recognition as sr
import PyPDF2
import docx
import os

# Configure Google Gemini API
genai.configure(api_key="AIzaSyBG-4wDu6xj90eMsoGHlrRp_yIG2xN5Rek")
model = genai.GenerativeModel("gemini-1.5-pro-latest")

# Speech-to-Text Function
def speech_to_text():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        recognizer.adjust_for_ambient_noise(source)
        print("Listening...")
        try:
            audio = recognizer.listen(source, timeout=5)
            return recognizer.recognize_google(audio)
        except sr.UnknownValueError:
            return "Could not understand the audio."
        except sr.RequestError:
            return "Speech recognition service error."
        except sr.WaitTimeoutError:
            return "No speech detected, try again."

# Document Processing Function
def process_file(file, pages=None):
    if not file:
        return "No document uploaded."
    
    ext = os.path.splitext(file.name)[-1].lower()
    doc_text = ""
    
    try:
        if ext == ".pdf":
            with open(file.name, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                if pages:
                    start_page, end_page = map(int, pages.split('-'))
                    doc_text = "\n".join([reader.pages[i - 1].extract_text() for i in range(start_page, end_page + 1)])
                else:
                    doc_text = "\n".join([page.extract_text() for page in reader.pages])

        elif ext in [".doc", ".docx"]:
            doc = docx.Document(file.name)
            doc_text = "\n".join([para.text for para in doc.paragraphs])

        elif ext == ".txt":
            with open(file.name, "r", encoding="utf-8") as f:
                doc_text = f.read()

        return doc_text if doc_text.strip() else "No readable text found in the document."
    
    except Exception as e:
        return f" Error processing document: {str(e)}"

#  Chatbot Processing
def chat_with_gemini(user_input, chat_history, image=None, file=None, pages=None):
    chat_history = chat_history or []
    combined_input = user_input.strip()

    if file:
        doc_text = process_file(file, pages)
        combined_input += f"\nExtracted Document Content:\n{doc_text}"

    if image:
        response = model.generate_content([combined_input, image])
    else:
        response = model.generate_content(combined_input)

    bot_response = response.text if response else "🤖 I'm sorry, I didn't understand that."

    chat_history.append(("👤 You", user_input))
    chat_history.append(("🤖 Gemini", bot_response))

    return chat_history, ""

#Improved UI
with gr.Blocks(theme=gr.themes.Default(primary_hue="blue")) as demo:
    gr.Markdown("<h1 style='text-align: center;'>🤖 Gemini AI Chatbot</h1>", elem_id="header")
    gr.Markdown("<h3 style='text-align: center;'>💬 Chat | 🎤 Voice | 📄 Documents | 🖼️ Images</h3>")

    with gr.Row():
        chatbox = gr.Chatbot(label="Conversation", height=400, elem_id="chatbox")

    with gr.Row():
        text_input = gr.Textbox(placeholder="Type your message...", label="Your Message", lines=2, elem_id="textbox")

    with gr.Row():
        voice_button = gr.Button("Speak", variant="primary")
        file_input = gr.File(label="Upload Document", file_types=[".pdf", ".docx", ".txt"])
        image_input = gr.Image(type="pil", label="Upload Image")
        page_range_input = gr.Textbox(label="Page Range (e.g., 3-5)", placeholder="Optional", visible=False)

    with gr.Row():
        send_button = gr.Button("Send", variant="primary")

    #Dynamic Visibility for Page Range Input
    def toggle_page_range_visibility(file):
        return gr.Textbox(visible=bool(file))

    file_input.change(fn=toggle_page_range_visibility, inputs=file_input, outputs=page_range_input)

    #Voice Button Updates Textbox
    voice_button.click(fn=speech_to_text, outputs=text_input)

    #Chat Submission
    send_button.click(
        fn=chat_with_gemini,
        inputs=[text_input, chatbox, image_input, file_input, page_range_input],
        outputs=[chatbox, text_input]
    )

demo.launch()


'''
import google.generativeai as genai
import gradio as gr
import speech_recognition as sr
import PyPDF2
import docx
from PIL import Image
import os

# Configure Google Gemini API
genai.configure(api_key="AIzaSyBG-4wDu6xj90eMsoGHlrRp_yIG2xN5Rek")
model = genai.GenerativeModel("gemini-1.5-pro-latest")

# 🎤 Improved Speech-to-Text Function
def speech_to_text():
    recognizer = sr.Recognizer()
    with sr.Microphone() as source:
        recognizer.adjust_for_ambient_noise(source)  # Reduce background noise
        print("Listening...")
        try:
            audio = recognizer.listen(source, timeout=5)  # 5-sec timeout to avoid long waits
            return recognizer.recognize_google(audio)
        except sr.UnknownValueError:
            return "❌ Could not understand the audio."
        except sr.RequestError:
            return "⚠️ Speech recognition service error."
        except sr.WaitTimeoutError:
            return "⏳ No speech detected, try again."

# 📄 Document Processing (PDF, DOCX, TXT)
def process_file(file, pages=None):
    if not file:
        return "⚠️ No document uploaded."

    ext = os.path.splitext(file.name)[-1].lower()
    doc_text = ""

    try:
        if ext == ".pdf":
            with open(file.name, "rb") as f:
                reader = PyPDF2.PdfReader(f)
                if pages:
                    start_page, end_page = map(int, pages.split('-'))
                    doc_text = "\n".join([reader.pages[i - 1].extract_text() for i in range(start_page, end_page + 1)])
                else:
                    doc_text = "\n".join([page.extract_text() for page in reader.pages])

        elif ext in [".doc", ".docx"]:
            doc = docx.Document(file.name)
            doc_text = "\n".join([para.text for para in doc.paragraphs])

        elif ext == ".txt":
            with open(file.name, "r", encoding="utf-8") as f:
                doc_text = f.read()

        return doc_text if doc_text.strip() else "⚠️ No readable text found in the document."
    
    except Exception as e:
        return f"❌ Error processing document: {str(e)}"

# 🤖 Chatbot Processing
def chat_with_gemini(user_input, chat_history, image=None, file=None, pages=None):
    chat_history = chat_history or []
    combined_input = user_input.strip()

    # Process document only if uploaded
    if file:
        doc_text = process_file(file, pages)
        combined_input += f"\n📄 Extracted Document Content:\n{doc_text}"

    # Process image & text together if an image is uploaded
    if image:
        response = model.generate_content([combined_input, image])
    else:
        response = model.generate_content(combined_input)

    bot_response = response.text if response else "🤖 I'm sorry, I didn't understand that."

    # Append responses to chat history
    chat_history.append(("👤 User", user_input))
    chat_history.append(("🤖 Bot", bot_response))

    return chat_history, ""

# 🎨 Improved UI Design
with gr.Blocks() as demo:
    gr.Markdown(
        "<h1 style='text-align: center;'>🤖 Gemini AI Chatbot</h1>"
        "<h3 style='text-align: center;'>💬 Chat | 🎤 Voice | 📄 Documents | 🖼️ Images</h3>"
    )

    

    with gr.Row():
        chatbox = gr.Chatbot(label="Chat History", height=350)

    with gr.Row():
        text_input = gr.Textbox(placeholder="Type your message here...", label="Your Message", lines=2)

    with gr.Row():
        voice_button = gr.Button("🎤 Speak", variant="primary")  # Voice input
        file_input = gr.File(label="📄 Upload Document (PDF, DOCX, TXT)", file_types=[".pdf", ".docx", ".txt"])
        image_input = gr.Image(type="pil", label="🖼️ Upload an Image")
        page_range_input = gr.Textbox(label="📑 Page Range (e.g., 3-5)", placeholder="Optional", visible=False)

    with gr.Row():
        send_button = gr.Button("🚀 Send", variant="primary")

    # 🛠️ Voice Button Fills Textbox
    voice_button.click(fn=speech_to_text, outputs=text_input)

    # 📨 Chat Submission (on button click)
    send_button.click(
        fn=chat_with_gemini,
        inputs=[text_input, chatbox, image_input, file_input, page_range_input],
        outputs=[chatbox, text_input]
    )

demo.launch()'
'''